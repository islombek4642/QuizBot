import json
import asyncio
import subprocess
import tempfile
import os
import base64
from typing import List, Dict, Optional, Tuple, Callable, Awaitable
from io import BytesIO
from core.config import settings
from core.logger import logger
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from groq import AsyncGroq, RateLimitError, APITimeoutError

class AIService:
    """Service for AI-powered quiz generation using Groq API."""
    
    def __init__(self):
        self.api_key = settings.GROQ_API_KEY
        self.model = settings.GROQ_MODEL
        self.client = AsyncGroq(
            api_key=self.api_key,
            max_retries=3
        )
        
    async def generate_quiz(self, topic: str, count: int = 30, lang: str = "UZ", 
                           on_progress: Optional[Callable[[int, int], Awaitable[None]]] = None) -> Tuple[List[Dict], Optional[str]]:
        """
        Generate quiz questions using Groq SDK with batching for large counts.
        """
        if not self.api_key:
            return [], "GROQ_API_KEY is not configured"
        
        all_questions = []
        batch_size = 15 # Generate 15 questions per batch for reliability
        
        # Build base system prompt
        if lang == "UZ":
            system_prompt = """Siz universitet darajasidagi professional professor va imtihon tuzuvchi ekspertsiz. 
Mavzuni chuqur tahlil qiling va talabalarni imtihonga tayyorlash uchun sifatli, Oliy ta'lim standartlariga mos testlar yarating.

SAVOL SIFATIGA QO'YILADIGAN TALABLAR:
1. **Bilim chuqurligi**: Shunchaki faktlarni eslab qolish emas (masalan, "Nima u?"), balki mantiqiy xulosa qilish, tahlil va amaliy qo'llashni talab qiladigan savollar tuzing.
2. **Murakkablik**: Savollar o'ylantiradigan, bir nechta faktni bog'laydigan bo'lsin. Bloom taksonomiyasining tahlil va sintez darajalaridan foydalaning.
3. **Aldamchi variantlar (Distractors)**: Noto'g'ri javoblar ham mantiqan to'g'riga o'xshash, mohiyatan xato, lekin chalg'ituvchi bo'lishi shart. Hazil yoki ochiq-oydin xato variantlar ishlatmang.
4. **Akademik til**: Rasmiy va ilmiy uslubda, imlo xatolarisiz yozing.

Javobni FAQAT quyidagi JSON formatida qaytaring:
{
  "questions": [
    {
      "question": "Mantiqan murakkab savol matni",
      "options": ["To'g'ri javob", "Chalg'ituvchi xato 1", "Chalg'ituvchi xato 2", "Chalg'ituvchi xato 3"],
      "correct_option_id": 0
    }
  ]
}

correct_option_id har doim 0 bo'lsin. Savol max 280 belgi, variantlar max 100 belgi."""
        else:
            system_prompt = """You are a professional university professor and examination expert. 
Analyze the topic deeply and create high-quality quiz questions that meet academic standards for higher education.

QUESTION QUALITY REQUIREMENTS:
1. **Depth of Knowledge**: Do not ask simple recall questions (e.g., "What is X?"). Focus on questions that require logical reasoning, analysis, and practical application.
2. **Complexity**: Questions should be thought-provoking and connect multiple concepts. Use higher-level Bloom's Taxonomy levels (Analysis, Evaluation).
3. **High-Quality Distractors**: Incorrect options must be plausible and logically related to the topic, yet factually incorrect. Avoid obvious or generic wrong answers.
4. **Academic Style**: Use formal, precise, and professional language.

Return ONLY the following JSON format:
{
  "questions": [
    {
      "question": "Logically complex question text",
      "options": ["Correct answer", "Plausible distractor 1", "Plausible distractor 2", "Plausible distractor 3"],
      "correct_option_id": 0
    }
  ]
}

correct_option_id should always be 0. Question max 280 chars, options max 100 chars."""

        current_count = 0
        attempts_without_progress = 0
        max_attempts = 10 
        
        while current_count < count and attempts_without_progress < max_attempts:
            remaining = count - current_count
            to_generate = min(batch_size, remaining)
            
            if lang == "UZ":
                user_prompt = f"Mavzu: {topic}\nSoni: {to_generate} ta yangi (takrorlanmagan) test savoli yarating."
            else:
                user_prompt = f"Topic: {topic}\nGenerate {to_generate} new (unique) quiz questions."
            
            try:
                response = await self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.8,
                    max_completion_tokens=4096,
                    # Pass service_tier if supported by installed version, otherwise via extra_body.
                    # extra_body is safe for both.
                    extra_body={"service_tier": settings.GROQ_SERVICE_TIER}
                )
                
                content = response.choices[0].message.content
                
                # Parse JSON
                batch_questions = []
                try:
                    batch_raw = json.loads(content)
                    if isinstance(batch_raw, dict) and "questions" in batch_raw:
                        batch_questions = batch_raw.get("questions", [])
                    elif isinstance(batch_raw, list): # Fallback
                        batch_questions = batch_raw
                except Exception as je:
                    logger.error("JSON parse failed in batch", error=str(je), content=content[:500])
                    attempts_without_progress += 1
                    continue
                
                # Validate and fix
                validated = self._validate_questions(batch_questions)
                
                if not validated:
                    logger.warning("Batch returned 0 valid questions", content=content[:500])
                    attempts_without_progress += 1
                    continue
                
                # Deduplication
                existing_questions = {q["question"] for q in all_questions}
                unique_validated = [q for q in validated if q["question"] not in existing_questions]
                
                if unique_validated:
                    all_questions.extend(unique_validated)
                    current_count = len(all_questions)
                    attempts_without_progress = 0
                    logger.info("Generation progress", topic=topic, current=current_count, total=count)
                    # Report progress
                    if on_progress:
                        await on_progress(min(current_count, count), count)
                else:
                    # Logic: We got valid JSON, but they were duplicates. 
                    # This is technically "bad progress", but maybe not a hard failure of the AI.
                    # We still count it as attempt_without_progress to prevent infinite loops if model keeps repeating.
                    logger.info("Batch generated duplicates only")
                    attempts_without_progress += 1
                    
                # Slow down slightly 
                if count > 50:
                    await asyncio.sleep(1)
                    
            except (RateLimitError, APITimeoutError) as e:
                logger.error(f"Groq API Error: {e}")
                if all_questions: break
                return [], f"Groq API Error: {str(e)}"
            except Exception as e:
                logger.exception(f"Batch generation error: {e}")
                if all_questions: break
                return [], f"Generation error: {str(e)}"
        
        if attempts_without_progress >= max_attempts:
            logger.error("AI generation stopped due to lack of progress", topic=topic, generated=len(all_questions))
            if not all_questions:
                return [], "AI failed to generate valid questions after multiple attempts"

        if not all_questions:
            return [], "Failed to generate any questions"
            
        logger.info("AI quiz generated", topic=topic, total=len(all_questions))
        return all_questions[:count], None

    async def convert_quiz(self, raw_text: str, lang: str = "UZ", on_progress: Optional[Callable[[int, int, int], Awaitable[None]]] = None) -> Tuple[List[Dict], Optional[str]]:
        """
        Convert raw text from PDF/Word to our quiz format using AI.
        Uses line-aware chunking and exhaustive parsing.
        """
        if not self.api_key:
            return [], "GROQ_API_KEY is not configured"

        # Line-aware chunking (approx 3500 chars to stay safe)
        max_chunk_chars = 3500
        lines = raw_text.splitlines()
        chunks = []
        current_chunk = []
        current_len = 0
        
        for line in lines:
            if current_len + len(line) > max_chunk_chars and current_chunk:
                chunks.append("\n".join(current_chunk))
                current_chunk = [line]
                current_len = len(line)
            else:
                current_chunk.append(line)
                current_len += len(line) + 1
        if current_chunk:
            chunks.append("\n".join(current_chunk))
            
        all_questions = []
        
        system_prompt = """You are a professional quiz extractor and educational content creator.
TASK: Extract ALL questions/topics from the provided text and convert them into a structured JSON quiz.

STRICT RULES:
1. EXHAUSTIVE EXTRACTION: Do not skip ANY question or topic found in the text. Every identifiable point must become a quiz question.
2. AUTO-FILL OPTIONS: If a question/topic has no options provided, CREATE 4 high-quality, academic-level options (1 correct + 3 plausible distractors).
3. JSON FORMAT: You MUST return a JSON object with a "questions" array.
4. If the source text explicitly marks the correct answer (examples: lines starting with '+' vs '=', or options prefixed with '#', or similar markers), you MUST use that marked option as the correct answer.
5. Regardless of source format, the returned JSON MUST place the correct answer at index 0 of the "options" array and set correct_option_id to 0.
5. LANGUAGE PRESERVATION: Use the SAME language as the input text (e.g., if input is Russian, output MUST be Russian). DO NOT translate to English or any other language.
6. DO NOT invent extra questions beyond what exists in the text. Avoid duplicates.

JSON STRUCTURE:
{
  "questions": [
    {
      "question": "Clear and concise question text",
      "options": ["Correct Answer", "Distractor 1", "Distractor 2", "Distractor 3"],
      "correct_option_id": 0
    }
  ]
}

CRITICAL: Return only the JSON object. Do not explain your work."""

        for i, chunk in enumerate(chunks):
            if not chunk.strip(): continue
            logger.info(f"Processing chunk {i+1}/{len(chunks)} ({len(chunk)} chars)")
            
            try:
                if i > 0:
                    await asyncio.sleep(1.5) # Rate limit protection

                response = await self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Convert the following text segment into JSON. Extract every possible question:\n\n{chunk}"}
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.1,
                    max_completion_tokens=4096,
                    extra_body={"service_tier": settings.GROQ_SERVICE_TIER}
                )
                
                content = response.choices[0].message.content
                chunk_questions = self._parse_response(content)
                
                if chunk_questions:
                    validated = self._validate_questions(chunk_questions)
                    # Dedupe across chunks (AI may repeat questions)
                    seen = {q.get("question", "").strip().lower() for q in all_questions if q.get("question")}
                    for q in validated:
                        key = q.get("question", "").strip().lower()
                        if not key or key in seen:
                            continue
                        all_questions.append(q)
                        seen.add(key)
                    
                if on_progress:
                    await on_progress(i + 1, len(chunks), len(all_questions))
                    
            except Exception as e:
                logger.error(f"Error in chunk {i+1}", error=str(e))
                continue

        if not all_questions:
            return [], "Fayldan hech qanday savol ajratib bo'lmadi."
            
        return all_questions, None
    
    def _parse_response(self, content: str) -> List[Dict]:
        """Parse JSON from AI response, handling potential formatting issues and truncation."""
        content = content.strip()
        
        # Helper to try parsing a string as JSON
        def try_parse(s):
            try:
                data = json.loads(s)
                # Handle {"questions": [...]} wrapper
                if isinstance(data, dict) and "questions" in data:
                    return data["questions"]
                return data
            except json.JSONDecodeError:
                # If it's a truncated array, try to close it
                if s.startswith('[') and not s.endswith(']'):
                    try:
                        # Try to find the last complete object
                        last_obj_end = s.rfind('}')
                        if last_obj_end != -1:
                            fixed = s[:last_obj_end + 1] + ']'
                            return json.loads(fixed)
                    except:
                        pass
                return None

        # 1. Try direct parse (or wrapper)
        parsed = try_parse(content)
        if parsed is not None and isinstance(parsed, list): return parsed
        
        # 2. Try to extract JSON from markdown code block
        if "```json" in content:
            start = content.find("```json") + 7
            end = content.find("```", start)
            candidate = content[start:end].strip() if end > start else content[start:].strip()
            parsed = try_parse(candidate)
            if parsed is not None and isinstance(parsed, list): return parsed

        # 3. Try to find JSON array
        if "[" in content:
            start = content.find("[")
            end = content.rfind("]") + 1
            candidate = content[start:end] if end > start else content[start:]
            parsed = try_parse(candidate)
            if parsed is not None and isinstance(parsed, list): return parsed
        
        # 4. Try to find JSON object (wrapper case where code block might be missing)
        if "{" in content:
            start = content.find("{")
            end = content.rfind("}") + 1
            if end > start:
                parsed = try_parse(content[start:end])
                if parsed is not None and isinstance(parsed, list): return parsed

        logger.error("Failed to parse AI response", content=content[:500])
        return []
    
    def _validate_questions(self, questions: List[Dict]) -> List[Dict]:
        """Validate and fix questions to meet requirements."""
        validated = []
        
        for q in questions:
            try:
                # Check required fields
                if "question" not in q or "options" not in q:
                    continue
                
                # Ensure exactly 4 options
                options = q["options"]
                if len(options) < 4:
                    continue
                if len(options) > 4:
                    options = options[:4]
                
                # Truncate if too long
                question_text = q["question"][:280] if len(q["question"]) > 280 else q["question"]
                options = [opt[:95] if len(opt) > 95 else opt for opt in options]
                
                # Ensure correct_option_id is valid
                # IMPORTANT: Based on prompt system "correct_option_id should always be 0". 
                # We enforce this at validation level to be safe.
                correct_id = 0
                
                validated.append({
                    "question": question_text,
                    "options": options,
                    "correct_option_id": correct_id
                })
                
            except Exception as e:
                logger.warning("Question validation failed", error=str(e))
                continue
        
        return validated
    
    async def close(self):
        """Close the AsyncGroq client session."""
        try:
            await self.client.close()
        except:
            pass


def _clean_xml_string(s: str) -> str:
    """Remove control characters and other strings that are not XML compatible."""
    if not s:
        return ""
    # Remove NULL bytes and other invalid XML control characters
    # Valid XML characters: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    import re
    # This regex matches characters NOT in the valid XML set
    illegal_xml_re = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1F\uD800-\uDFFF\uFFFE\uFFFF]')
    return illegal_xml_re.sub('', s)


def generate_docx_from_questions(questions: List[Dict], title: str) -> bytes:
    """
    Generate a .docx file from questions in our format.
    """
    from docx import Document
    from docx.shared import Pt
    from io import BytesIO
    
    doc = Document()
    
    # Add title
    doc.add_heading(_clean_xml_string(title), 0)
    
    # Add each question in our format
    for i, q in enumerate(questions, 1):
        # Add question
        doc.add_paragraph(f"?{_clean_xml_string(q['question'])}")
        
        # Add options
        correct_id = q['correct_option_id']
        for j, opt in enumerate(q['options']):
            cleaned_opt = _clean_xml_string(opt)
            if j == correct_id:
                doc.add_paragraph(f"+{cleaned_opt}")
            else:
                doc.add_paragraph(f"={cleaned_opt}")
        
        # Add empty line between questions
        doc.add_paragraph()
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()
    logger.info("Generated DOCX", questions_count=len(questions), size_bytes=len(content))
    return content


async def extract_text_from_pdf(pdf_bytes: bytes, on_progress: Optional[Callable] = None) -> str:
    """Extract text from PDF using PyMuPDF with Vision OCR fallback."""
    # Check signature: %PDF-
    if not pdf_bytes.startswith(b'%PDF-'):
        logger.error("Invalid PDF signature")
        return ""

    text = ""
    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            page_count = len(doc)
            logger.info("PDF opened", pages=page_count, size=len(pdf_bytes))
            
            # Try normal extraction first (Optimized with join)
            texts = []
            for page in doc:
                texts.append(page.get_text())
            text = "".join(texts)
            
            if not text.strip() and page_count > 0:
                logger.warning("No text extracted from PDF, initiating Groq Vision OCR fallback", pages=page_count)
                text = await _extract_text_via_vision(doc, on_progress)
                    
    except Exception as e:
        logger.error("PDF extraction failed", error=str(e))
    return text

async def _extract_text_via_vision(doc: fitz.Document, on_progress: Optional[Callable] = None) -> str:
    """Uses Groq Vision to perform OCR on PDF pages via AsyncGroq SDK."""
    full_text = ""
    total_pages = len(doc)
    
    # Instantiate client just for OCR
    ocr_client = AsyncGroq(
        api_key=settings.GROQ_API_KEY,
        max_retries=3
    )

    try:
        for i, page in enumerate(doc):
            try:
                if on_progress:
                    if asyncio.iscoroutinefunction(on_progress):
                        await on_progress(i + 1, total_pages)
                
                # Render page to image (JPEG)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # 2x zoom
                img_bytes = pix.tobytes("jpeg")
                base64_image = base64.b64encode(img_bytes).decode('utf-8')
                
                response = await ocr_client.chat.completions.create(
                    model=settings.GROQ_VISION_MODEL,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "Extract all text from this image as plain text. Do not add any comments or explanations. Just the text content."},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{base64_image}",
                                    },
                                },
                            ],
                        }
                    ],
                    temperature=0.1,
                    extra_body={"service_tier": settings.GROQ_SERVICE_TIER}
                )
                
                if response.choices and response.choices[0].message.content:
                    page_text = response.choices[0].message.content
                    full_text += page_text + "\n\n"
                    logger.debug("Page OCR success", page=i+1)
                else:
                    logger.error("Groq Vision API returned empty content")
                
                # Cooldown to avoid hitting rate limits too fast
                await asyncio.sleep(0.5)
                
            except Exception as e:
                logger.error(f"OCR failed for page {i+1}", error=str(e))
                continue
    finally:
        await ocr_client.close()
            
    return full_text


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from Word document including tables using python-docx."""
    text = ""
    try:
        from io import BytesIO
        doc = DocxDocument(BytesIO(docx_bytes))
        
        # 1. Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # 2. Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"
                    
    except Exception as e:
        logger.error("DOCX extraction failed", error=str(e))
    return text


def extract_text_from_doc(doc_bytes: bytes) -> str:
    """
    Robust extraction for .doc files.
    """
    import subprocess
    import tempfile
    import os

    # 1. Check for RTF signature
    if doc_bytes.startswith(b'{\\rtf'):
        logger.info("Detected RTF format, using striprtf")
        try:
            from striprtf.striprtf import rtf_to_text
            return rtf_to_text(doc_bytes.decode('utf-8', errors='ignore'))
        except Exception as e:
            logger.error("RTF extraction failed", error=str(e))

    # 2. Check for DOCX (ZIP) signature
    if doc_bytes.startswith(b'PK\x03\x04'):
        logger.info("Detected DOCX format renamed to .doc, using docx parser")
        try:
            return extract_text_from_docx(doc_bytes)
        except Exception as e:
            logger.error("Docx fallback failed", error=str(e))

    # 3. Handle as legacy Word (.doc)
    text = ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
        tmp.write(doc_bytes)
        tmp_path = tmp.name

    try:
        # Try Antiword first
        process = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore'
        )
        if process.returncode == 0 and process.stdout.strip():
            text = process.stdout
        else:
            # Try Catdoc as fallback
            logger.info("Antiword failed or empty, trying catdoc")
            process_cat = subprocess.run(
                ["catdoc", "-w", tmp_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            if process_cat.returncode == 0:
                text = process_cat.stdout
            else:
                logger.error("Catdoc failed", stderr=process_cat.stderr)
                
    except Exception as e:
        logger.error("Legacy Word extraction failed", error=str(e))
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
            
    return text
