import docx
from typing import List, Dict, Any
import io

def generate_quiz_docx(title: str, questions: List[Dict[str, Any]]) -> io.BytesIO:
    """Generates a .docx file from a quiz's questions."""
    doc = docx.Document()
    doc.add_heading(title, 0)
    
    for i, q in enumerate(questions, 1):
        # Question text
        para = doc.add_paragraph(style='List Number')
        para.add_run(q['question']).bold = True
        
        # Options
        for j, opt in enumerate(q['options']):
            prefix = "+" if j == q['correct_option_id'] else "="
            doc.add_paragraph(f"{prefix} {opt}")
        
        # Explanation (if any)
        if q.get('explanation'):
            doc.add_paragraph(f"! {q['explanation']}", style='Caption')
            
        doc.add_paragraph() # Spacer
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
